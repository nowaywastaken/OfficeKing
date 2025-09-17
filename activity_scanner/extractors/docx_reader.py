from __future__ import annotations

"""DOCX extraction utilities."""

from typing import List


def read_docx_text(path: str) -> str:
    """Extract plain text from document paragraphs and tables."""

    from docx import Document  # type: ignore[import-not-found]

    document = Document(path)
    pieces: List[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    pieces.append(cell.text)
    return "\n".join(pieces)
