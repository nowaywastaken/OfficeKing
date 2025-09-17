from __future__ import annotations

from typing import List


def extract_text_from_docx(path: str) -> str:
    """Extract plain text from document paragraphs and tables."""
    from docx import Document  # type: ignore[import-not-found]

    document = Document(path)
    texts: List[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)
    return "\n".join(texts)
